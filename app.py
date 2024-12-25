import streamlit as st
from fpdf import FPDF
from docx import Document

# DOCX shablonni yuklash funksiyasi
def load_test_from_docx(file_path):
    doc = Document(file_path)
    content = []
    for paragraph in doc.paragraphs:
        content.append(paragraph.text.strip())
    content = "\n".join(content)
    
    questions = []
    question_blocks = content.strip().split("%%%%")
    for block in question_blocks:
        if block.strip():
            parts = block.strip().split("++++")
            question_text = parts[0].replace("****[1]\n", "").strip()
            options = [part.strip() for part in parts[1:]]
            correct_answer = options[0]
            questions.append({
                "question": question_text,
                "options": options,
                "correct_answer": correct_answer
            })
    return questions

# Savollarni bo‘limlarga ajratish
def split_questions(questions, chunk_size=25):
    return [questions[i:i + chunk_size] for i in range(0, len(questions), chunk_size)]

# Natijalarni hisoblash funksiyasi
def calculate_score(questions, user_answers):
    correct_count = 0
    for i, question in enumerate(questions):
        if question["correct_answer"] == user_answers.get(str(i), ""):
            correct_count += 1
    return correct_count

# Unicode shrift qo‘llab-quvvatlovchi PDF hisobotni yaratish funksiyasi
def generate_pdf_report(questions, user_answers, score):
    pdf = FPDF()
    pdf.add_page()
    pdf.add_font("ArialUnicode", "", "ARIAL.TTF", uni=True)
    pdf.set_font("ArialUnicode", size=12)
    pdf.cell(200, 10, txt="Test Natijalari", ln=True, align='C')
    pdf.ln(10)

    for i, question in enumerate(questions):
        pdf.set_font("ArialUnicode", size=10)
        pdf.multi_cell(0, 10, txt=f"{i + 1}. {question['question']}")
        pdf.multi_cell(0, 10, txt=f"  To'g'ri javob: {question['correct_answer']}")
        pdf.multi_cell(0, 10, txt=f"  Sizning javobingiz: {user_answers.get(str(i), '')}")
        pdf.ln(5)

    pdf.set_font("ArialUnicode", size=12)
    pdf.cell(0, 10, txt=f"To'g'ri javoblar soni: {score}/{len(questions)}", ln=True)
    return pdf

# Streamlit interfeysi
def main():
    st.markdown(
        """
        <style>
        .main-title {
            color: #4CAF50;
            text-align: center;
            font-family: 'Arial', sans-serif;
            font-size: 2.5em;
        }
        .subtitle {
            color: #FF5722;
            text-align: center;
            font-family: 'Courier New', Courier, monospace;
            font-size: 1.5em;
        }
        .question {
            background-color: #E0F7FA;
            padding: 15px;
            border-radius: 10px;
            margin-bottom: 10px;
        }
        .option {
            display: flex;
            align-items: center;
            background-color: #F1F8E9;
            padding: 10px;
            margin: 5px;
            border: 1px solid #CDDC39;
            border-radius: 5px;
            font-size: 16px;
            color: #4CAF50;
            cursor: pointer;
        }
        .option input {
            margin-right: 10px;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )

    st.markdown("<h1 class='main-title'>📘 Streamlit Test Tizimi</h1>", unsafe_allow_html=True)
    st.markdown("<h3 class='subtitle'>Quyidagi testga javob bering:</h3>", unsafe_allow_html=True)

    uploaded_file = st.file_uploader("Shablonni yuklang (DOCX format):", type=["docx"])
    
    if uploaded_file:
        questions = load_test_from_docx(uploaded_file)
        chunks = split_questions(questions, chunk_size=25)
        chunk_titles = [f"Savollar {i * 25 + 1}-{(i + 1) * 25}" for i in range(len(chunks))]
        
        selected_chunk_index = st.selectbox("Savollar bo‘limini tanlang:", options=range(len(chunks)), format_func=lambda x: chunk_titles[x])
        selected_questions = chunks[selected_chunk_index]

        font_size = st.slider("Shrift o‘lchamini tanlang:", min_value=12, max_value=44, value=16)

        user_answers = {}

        for i, question in enumerate(selected_questions):
            st.markdown(f"<div class='question' style='font-size:{font_size}px;'><b>{i + 1}. {question['question']}</b></div>", unsafe_allow_html=True)
            
            options_html = "".join([
                f"""
                <div class='option'>
                    <input type='radio' id='q{i}_{j}' name='q{i}' value='{option}' onchange="fetch('/api/submit', {{
                        method: 'POST',
                        body: JSON.stringify({{'question_id': {i}, 'answer': '{option}'}})
                    }})">
                    <label for='q{i}_{j}'>{option}</label>
                </div>
                """
                for j, option in enumerate(question["options"])
            ])
            st.markdown(options_html, unsafe_allow_html=True)
            user_answers[str(i)] = st.session_state.get(f"q{i}", "")

        if st.button("Testni yakunlash"):
            score = calculate_score(selected_questions, user_answers)
            st.success(f"✅ Test yakunlandi! To'g'ri javoblar soni: {score}/{len(selected_questions)}")
            pdf = generate_pdf_report(selected_questions, user_answers, score)
            pdf_file_path = "test_results.pdf"
            pdf.output(pdf_file_path)
            
            with open(pdf_file_path, "rb") as pdf_file:
                pdf_data = pdf_file.read()
            st.download_button("📘 Natijalarni PDF shaklda yuklab olish", data=pdf_data, file_name="test_results.pdf", mime="application/pdf")

if __name__ == "__main__":
    main()


# import streamlit as st
# from fpdf import FPDF
# from docx import Document

# # DOCX shablonni yuklash funksiyasi
# def load_test_from_docx(file_path):
#     doc = Document(file_path)
#     content = []
#     for paragraph in doc.paragraphs:
#         content.append(paragraph.text.strip())
#     content = "\n".join(content)
    
#     questions = []
#     question_blocks = content.strip().split("%%%%")
#     for block in question_blocks:
#         if block.strip():
#             parts = block.strip().split("++++")
#             question_text = parts[0].replace("****[1]\n", "").strip()
#             options = [part.strip() for part in parts[1:]]
#             correct_answer = options[0]
#             questions.append({
#                 "question": question_text,
#                 "options": options,
#                 "correct_answer": correct_answer
#             })
#     return questions

# # Savollarni bo‘limlarga ajratish
# def split_questions(questions, chunk_size=25):
#     return [questions[i:i + chunk_size] for i in range(0, len(questions), chunk_size)]

# # Natijalarni hisoblash funksiyasi
# def calculate_score(questions, user_answers):
#     correct_count = 0
#     for i, question in enumerate(questions):
#         if question["correct_answer"] == user_answers[i]:
#             correct_count += 1
#     return correct_count

# # Unicode shrift qo‘llab-quvvatlovchi PDF hisobotni yaratish funksiyasi
# def generate_pdf_report(questions, user_answers, score):
#     pdf = FPDF()
#     pdf.add_page()
#     pdf.add_font("ArialUnicode", "", "ARIAL.TTF", uni=True)
#     pdf.set_font("ArialUnicode", size=12)
#     pdf.cell(200, 10, txt="Test Natijalari", ln=True, align='C')
#     pdf.ln(10)

#     for i, question in enumerate(questions):
#         pdf.set_font("ArialUnicode", size=10)
#         pdf.multi_cell(0, 10, txt=f"{i + 1}. {question['question']}")
#         pdf.multi_cell(0, 10, txt=f"  To'g'ri javob: {question['correct_answer']}")
#         pdf.multi_cell(0, 10, txt=f"  Sizning javobingiz: {user_answers[i]}")
#         pdf.ln(5)

#     pdf.set_font("ArialUnicode", size=12)
#     pdf.cell(0, 10, txt=f"To'g'ri javoblar soni: {score}/{len(questions)}", ln=True)
#     return pdf

# # Streamlit interfeysi
# def main():
#     st.markdown(
#         """
#         <style>
#         .main-title {
#             color: #4CAF50;
#             text-align: center;
#             font-family: 'Arial', sans-serif;
#             font-size: 2.5em;
#         }
#         .subtitle {
#             color: #FF5722;
#             text-align: center;
#             font-family: 'Courier New', Courier, monospace;
#             font-size: 1.5em;
#         }
#         .question {
#             background-color: #E0F7FA;
#             padding: 15px;
#             border-radius: 10px;
#             margin-bottom: 10px;
#         }
#         .option {
#             display: flex;
#             align-items: center;
#             background-color: #F1F8E9;
#             padding: 10px;
#             margin: 5px;
#             border: 1px solid #CDDC39;
#             border-radius: 5px;
#             font-size: 16px;
#             color: #4CAF50;
#         }
#         .radio-button {
#             margin-right: 10px;
#         }
#         </style>
#         """,
#         unsafe_allow_html=True,
#     )

#     st.markdown("<h1 class='main-title'>📘 Streamlit Test Tizimi</h1>", unsafe_allow_html=True)
#     st.markdown("<h3 class='subtitle'>Quyidagi testga javob bering:</h3>", unsafe_allow_html=True)

#     uploaded_file = st.file_uploader("Shablonni yuklang (DOCX format):", type=["docx"])
    
#     if uploaded_file:
#         questions = load_test_from_docx(uploaded_file)
#         chunks = split_questions(questions, chunk_size=25)
#         chunk_titles = [f"Savollar {i * 25 + 1}-{(i + 1) * 25}" for i in range(len(chunks))]
        
#         selected_chunk_index = st.selectbox("Savollar bo‘limini tanlang:", options=range(len(chunks)), format_func=lambda x: chunk_titles[x])
#         selected_questions = chunks[selected_chunk_index]

#         font_size = st.slider("Shrift o‘lchamini tanlang:", min_value=12, max_value=44, value=16)

#         user_answers = []
#         progress = 0
#         for i, question in enumerate(selected_questions):
#             st.markdown(f"<div class='question' style='font-size:{font_size}px;'><b>{i + 1}. {question['question']}</b></div>", unsafe_allow_html=True)
            
#             selected_option = st.radio(
#                 label="Javobingizni tanlang:",
#                 options=question["options"],
#                 key=f"q{selected_chunk_index}_{i}",
#                 format_func=lambda x: x,  # Disable formatting for HTML
#                 horizontal=False,
#             )

#             options_html = "".join([
#                 f"<div class='option'><input class='radio-button' type='radio' name='q{i}' value='{option}' disabled {'checked' if option == selected_option else ''} />{option}</div>"
#                 for option in question["options"]
#             ])
#             st.markdown(options_html, unsafe_allow_html=True)

#             user_answers.append(selected_option)
#             progress += 1
#             st.progress(progress / len(selected_questions))

#         if st.button("Testni yakunlash"):
#             score = calculate_score(selected_questions, user_answers)
#             st.success(f"✅ Test yakunlandi! To'g'ri javoblar soni: {score}/{len(selected_questions)}")
#             st.markdown("<h3 class='subtitle'>Natijalar:</h3>", unsafe_allow_html=True)
#             for i, question in enumerate(selected_questions):
#                 st.markdown(f"<div class='question'><b>{i + 1}. {question['question']}</b></div>", unsafe_allow_html=True)
#                 st.markdown(f"<span class='correct-answer'>To'g'ri javob: {question['correct_answer']}</span>", unsafe_allow_html=True)
#                 st.markdown(f"<span class='user-answer'>Sizning javobingiz: {user_answers[i]}</span>", unsafe_allow_html=True)

#             pdf = generate_pdf_report(selected_questions, user_answers, score)
#             pdf_file_path = "test_results.pdf"
#             pdf.output(pdf_file_path)
            
#             with open(pdf_file_path, "rb") as pdf_file:
#                 pdf_data = pdf_file.read()
#             st.download_button("📘 Natijalarni PDF shaklda yuklab olish", data=pdf_data, file_name="test_results.pdf", mime="application/pdf")

# if __name__ == "__main__":
#     main()

